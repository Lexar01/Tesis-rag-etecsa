# %%
import json
import numpy as np
from tensorflow import keras
import os
import re
import unicodedata
from pypdf import PdfReader
from docx import Document
import psycopg2
import logging
from datetime import datetime
import io

# %%

 
logging.basicConfig(
    filename="system.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

def log_event(event_type, message):
   
    logging.info(f"[{event_type}] {message}")


# %%
try:
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
except NameError:
     
    BASE_DIR = os.path.dirname(os.getcwd())

# %%
ENC_SAVE_PATH = os.path.join(BASE_DIR, "storage", "Archivos json de los codigos", "enc_definitivo.json")
VOCAB_SAVE_PATH = os.path.join(BASE_DIR, "storage", "Archivos json de los codigos", "vocabulary_definitivo.json")
EMB_OUT = os.path.join(BASE_DIR, "storage", "Archivos json de los codigos", "embeddings_final.npy")
FOLDER = os.path.join(BASE_DIR, "storage", "Dataset para convertir en vectores")

# %%
def normalize_text(text):
     
    text = text.lower()
    text = ''.join(
        c for c in unicodedata.normalize('NFD', text)
        if unicodedata.category(c) != 'Mn'
    )
    return text

# %%
def extract_text_from_pdf(pdf_path):
    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        text_page = page.extract_text()
        if text_page:
            text += text_page + "\n"
    return text

def extract_text_from_docx(docx_path_or_bytes):
    if isinstance(docx_path_or_bytes, (bytes, bytearray)):
        stream = io.BytesIO(docx_path_or_bytes)
        doc = Document(stream)
    elif hasattr(docx_path_or_bytes, "read"):
        try:
            docx_path_or_bytes.seek(0)
        except Exception:
            pass
        doc = Document(docx_path_or_bytes)
    else:
        doc = Document(docx_path_or_bytes)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def extract_text_from_txt(txt_path):
    with open(txt_path, "r", encoding="utf-8") as f:
        return f.read()

# %%
def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def chunk_text(text, chunk_size=512):
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size):
        chunk = " ".join(words[i:i+chunk_size])
        chunks.append(chunk)
    return chunks

# %%
conn = psycopg2.connect(
    dbname="Tesis",
    user="postgres",
    password="1234",
    host="localhost",
    port="5432"
)
cursor = conn.cursor()

# %%
def ingest_documents(folder_path, chunk_size=512):
    corpus = []
    if not os.path.exists(folder_path) or not os.listdir(folder_path):
        return corpus

    for filename in os.listdir(folder_path):
        filepath = os.path.join(folder_path, filename)
        try:
            if filename.lower().endswith(".pdf"):
                raw_text = extract_text_from_pdf(filepath)
            elif filename.lower().endswith(".txt"):
                raw_text = extract_text_from_txt(filepath)
            elif filename.lower().endswith(".docx"):
                raw_text = extract_text_from_docx(filepath)
            else:
                continue
        except Exception as e:
            print(f"Error al procesar {filename}: {e}")
            continue
        cleaned_text = clean_text(raw_text)
        chunks = chunk_text(cleaned_text, chunk_size)

        for idx, chunk in enumerate(chunks, start=1):
            corpus.append({
                "chunk_id": idx,
                "text": chunk,
                "metadata": {"source": filename}
            })
    return corpus

# %%


emb_matrix = np.load(EMB_OUT)

with open(ENC_SAVE_PATH,"r", encoding="utf-8") as f:   
    enc = json.load(f)

with open(VOCAB_SAVE_PATH,"r", encoding="utf-8") as f:   
    vocabulary = json.load(f)

"""
enc
"""

# %%
norms = np.linalg.norm(emb_matrix, axis=1, keepdims=True)
norms[norms == 0] = 1.0
emb_matrix = emb_matrix / norms

# %%
def get_vector(word):
    idx = enc.get(word)
    return emb_matrix[idx] if idx is not None else None

# %%
def chunk_embedding(chunk, dim=200):
    vectors = []
     
    chunk = normalize_text(chunk)
    for word in chunk.split():
        vec = get_vector(word)
        if vec is not None:
            vectors.append(vec)
    if vectors:
        return np.mean(vectors, axis=0)
    else:
        return np.zeros(dim)

# %%
def generate_embeddings(corpus, dim=200):
    results = []
    for item in corpus:
        norm_text = normalize_text(item["text"])
        vector = chunk_embedding(norm_text, dim)
        results.append({
            "chunk_id": item["chunk_id"],
            "embedding": vector.tolist(),
            "metadata": item["metadata"],
            "text": item["text"]  
        })
    return results

# %%
def insert_embedding(chunk_id, embedding, metadata, text, almacenamiento="permanente"):
    try:
        if np.all(embedding == 0):
            return
        cursor.execute("""
            SELECT 1 FROM chunk WHERE embedding = %s
        """, (json.dumps(embedding.tolist()),))
        if cursor.fetchone() is None:
            cursor.execute("""
                INSERT INTO chunk (chunk_id, embedding, metadata, text, almacenamiento)
                VALUES (%s, %s, %s, %s, %s)
            """, (
                chunk_id,
                json.dumps(embedding.tolist()),
                json.dumps(metadata),
                text,
                almacenamiento
            ))
            conn.commit()
    except Exception as e:
        print("Error al insertar:", e)
        conn.rollback()

# %%
def get_next_id(table_name: str, id_column: str = "id"):
    cursor.execute(f"""
        SELECT MIN(t1.{id_column} + 1)
        FROM {table_name} t1
        LEFT JOIN {table_name} t2
          ON t2.{id_column} = t1.{id_column} + 1
        WHERE t2.{id_column} IS NULL
    """)
    result = cursor.fetchone()[0]
    if result is None:
        cursor.execute(f"SELECT COALESCE(MAX({id_column}),0)+1 FROM {table_name}")
        result = cursor.fetchone()[0]
    return result

# %%
def populate_embeddings(embeddings):
    for item in embeddings:
        insert_embedding(
            get_next_id("chunk", "chunk_id"),
            np.array(item["embedding"], dtype="float32"),
            item["metadata"],
            item["text"]
        )

# %%
def get_all_embeddings():
    cursor.execute("SELECT chunk_id, embedding, metadata, text FROM chunk")
    rows = cursor.fetchall()
    data = []
    for chunk_id, emb, meta, text in rows:
        if isinstance(emb, str):
            emb = np.array(json.loads(emb), dtype="float32")
        elif isinstance(emb, list):
            emb = np.array(emb, dtype="float32")
        data.append({
            "chunk_id": chunk_id,
            "embedding": emb,
            "metadata": json.loads(meta) if isinstance(meta, str) else meta,
            "text": text 
        })
    return data


# %%
def cosine_similarity(a, b):
    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))

def contains_word(text, word):
     
    return re.search(rf"\b{word}\b", text) is not None


# %%
def retrieve(query_text, k=5):
    
    query_norm = normalize_text(query_text)
    query_vec = chunk_embedding(query_norm)

    registros = get_all_embeddings()
    scores = []
    for r in registros:
 
        text_norm = normalize_text(r["text"])
        sim = cosine_similarity(query_vec, r["embedding"])

         
        bonus = 0.0
        for word in query_norm.split():
            if contains_word(text_norm, word):
                bonus += 0.0   
         
        score = sim + bonus
        scores.append((score, r))
    
    log_event("RETRIEVE", f"Query='{query_text}' | Top score={scores[0][0]:.2f}")
     
    scores.sort(reverse=True, key=lambda x: x[0])
    return scores[:k]


# %%
def database():
    try:
        data = ingest_documents(FOLDER, chunk_size=512)
        if not data:
            log_event("DATABASE_WARNING", "No se encontraron documentos para procesar.")
            return {"docs": 0, "embeddings": 0, "status": "warning"}

        embeddings = generate_embeddings(data)
        populate_embeddings(embeddings)

        log_event("DATABASE", f"Pipeline completado. Documentos={len(data)}, Embeddings={len(embeddings)}")
        return {"docs": len(data), "embeddings": len(embeddings), "status": "ok"}

    except Exception as e:
        log_event("DATABASE_ERROR", f"Error general en database(): {e}")
        return {"docs": 0, "embeddings": 0, "status": "error"}

